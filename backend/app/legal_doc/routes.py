from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional
import html
import os
from PyPDF2 import PdfReader
import re
import uuid
from pathlib import Path
from datetime import datetime

# LEGALAI: безопасный импорт python-docx.
# В Termux python-docx может отсутствовать, на сервере Ubuntu он будет установлен.
try:
    from docx import Document
except ImportError:
    Document = None

from app.cases_db import get_db
from app.auth.utils import get_current_user
from app.auth.models import User
from .models import (
    Template as TemplateModel,
    Document as DocumentModel,
    Case as CaseModel,
    Attachment as AttachmentModel,
)
from .schemas import (
    TemplateOut,
    DocumentOut,
    DocumentListItem,
    DocumentCreate,
    CaseCreate,
    CaseRead,
    WorkspaceDocumentCreate,
    WorkspaceDocumentUpdate,
    WorkspaceDocumentRead,
    AttachmentRead,
)
from .utils import save_template_file, fill_template
from .pdf_utils import convert_docx_to_pdf


# LEGALAI: АКТИВНЫЕ РОУТЫ ДЛЯ ШАБЛОНОВ, ДОКУМЕНТОВ, ДЕЛ И WORKSPACE.


router = APIRouter()


def extract_fields_from_docx(path: str) -> List[str]:
    """Извлечь плейсхолдеры вида {{...}} из DOCX."""
    # Если python-docx недоступен (Termux), просто возвращаем пустой список,
    # чтобы не ломать остальное API.
    if Document is None:
        return []

    placeholders: List[str] = []
    try:
        doc = Document(path)
    except Exception:
        return placeholders
    text = ""
    for paragraph in doc.paragraphs:
        text += " " + paragraph.text
    # Find all occurrences of placeholders {{...}}
    matches = re.findall(r"{{(.*?)}}", text)
    # Remove duplicates while preserving order
    seen = set()
    unique_fields: List[str] = []
    for match in matches:
        field = match.strip()
        if field not in seen:
            seen.add(field)
            unique_fields.append(field)
    return unique_fields


# -------- TEMPLATES (ШАБЛОНЫ) --------


@router.post("/templates", response_model=TemplateOut)
async def upload_template(
    name: str = Form(...),
    description: Optional[str] = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Загрузить новый шаблон .docx и создать запись в БД."""
    # Validate file extension
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Only .docx files allowed",
        )

    # Validate file size (<= 5MB)
    try:
        current_pos = file.file.tell()
        file.file.seek(0, os.SEEK_END)
        size = file.file.tell()
        file.file.seek(current_pos, os.SEEK_SET)
    except Exception:
        size = 0
    if size > 5 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File too large",
        )

    # Validate that it's a proper docx file
    # В Termux, где python-docx недоступен, этот шаг пропускаем.
    if Document is not None:
        try:
            Document(file.file)
            file.file.seek(0)
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid .docx format",
            )
    else:
        # На всякий случай вернём файловый указатель в начало.
        file.file.seek(0)

    file_path = await save_template_file(file, user.id)
    template = TemplateModel(
        name=name,
        description=description,
        file_path=file_path,
        owner_id=user.id,
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    return template


@router.get("/templates", response_model=List[TemplateOut])
async def list_templates(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Список шаблонов текущего пользователя."""
    templates = (
        db.query(TemplateModel)
        .filter(TemplateModel.owner_id == user.id)
        .all()
    )
    return templates


@router.get("/templates/{id}/fields", response_model=List[str])
async def get_template_fields(
    id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Извлечь плейсхолдеры из шаблона."""
    template = (
        db.query(TemplateModel)
        .filter(
            TemplateModel.id == id,
            TemplateModel.owner_id == user.id,
        )
        .first()
    )
    if not template:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
    fields = extract_fields_from_docx(template.file_path)
    return fields


# -------- DOCUMENTS (DOCX + PDF из шаблонов) --------


@router.post("/documents", response_model=DocumentOut)
async def generate_document(
    document_data: DocumentCreate,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Сгенерировать документ из шаблона по переданным данным."""
    # В среде без python-docx (Termux) генерация DOCX недоступна.
    if Document is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX generation is unavailable in this environment. "
                   "Запусти backend на сервере, где установлен python-docx.",
        )

    template = (
        db.query(TemplateModel)
        .filter(
            TemplateModel.id == document_data.template_id,
            TemplateModel.owner_id == user.id,
        )
        .first()
    )
    if not template:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")

    # Ensure provided values correspond to placeholders in the template
    allowed_fields = extract_fields_from_docx(template.file_path)
    for key in document_data.values.keys():
        if key not in allowed_fields:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid field: {key}",
            )

    # Escape HTML to prevent XSS/injection
    cleaned_values = {k: html.escape(v) for k, v in document_data.values.items()}
    file_path = fill_template(template.file_path, cleaned_values)

    # Optionally generate PDF
    if getattr(document_data, "generate_pdf", False):
        try:
            convert_docx_to_pdf(file_path)
        except Exception:
            # If conversion fails, we still proceed with the DOCX
            pass

    name = os.path.basename(file_path)
    document = DocumentModel(
        template_id=template.id,
        user_id=user.id,
        file_path=file_path,
        status="completed",  # документы из шаблона считаем завершёнными
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    return DocumentOut(
        id=document.id,
        name=name,
        file_path=document.file_path,
        created_at=document.created_at,
    )


@router.get("/documents", response_model=List[DocumentListItem])
async def list_documents(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Список документов (DOCX) пользователя, созданных из шаблонов."""
    documents = (
        db.query(DocumentModel)
        .filter(DocumentModel.user_id == user.id)
        .order_by(DocumentModel.created_at.desc())
        .all()
    )
    items: List[DocumentListItem] = []
    for d in documents:
        name = os.path.basename(d.file_path)
        items.append(
            DocumentListItem(
                id=d.id,
                name=name,
                created_at=d.created_at,
            )
        )
    return items


@router.get("/documents/{id}")
async def get_document(
    id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Скачать сгенерированный DOCX-документ по ID."""
    document = (
        db.query(DocumentModel)
        .filter(
            DocumentModel.id == id,
            DocumentModel.user_id == user.id,
        )
        .first()
    )
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    filename = os.path.basename(document.file_path)
    return FileResponse(
        path=document.file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/documents/{id}/pdf")
async def get_document_pdf(
    id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Скачать PDF-версию документа. Если её нет — сгенерировать из DOCX."""
    document = (
        db.query(DocumentModel)
        .filter(
            DocumentModel.id == id,
            DocumentModel.user_id == user.id,
        )
        .first()
    )
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    pdf_path = document.file_path.replace(".docx", ".pdf")
    if not os.path.exists(pdf_path):
        try:
            convert_docx_to_pdf(document.file_path)
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=str(exc),
            )
    filename = os.path.basename(pdf_path)
    return FileResponse(path=pdf_path, filename=filename, media_type="application/pdf")


# -------- WORKSPACE: CASES / DOCUMENTS / ATTACHMENTS --------


@router.post("/workspace/cases/start", response_model=CaseRead)
def start_case(
    payload: CaseCreate,
    db: Session = Depends(get_db),
):
    """
    Создать новое дело.
    В будущем user_id должен браться из токена, сейчас приходит в payload.
    """
    title = payload.title or f"Вопрос от {datetime.utcnow().strftime('%Y-%m-%d')}"
    new_case = CaseModel(
        user_id=payload.user_id,
        title=title,
        status="active",
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
        last_user_activity_at=datetime.utcnow(),
    )
    db.add(new_case)
    db.commit()
    db.refresh(new_case)
    return new_case


@router.get("/workspace/cases", response_model=List[CaseRead])
def list_cases(
    user_id: int,
    db: Session = Depends(get_db),
):
    """
    Список дел пользователя для блока 'Мои дела'.
    """
    cases = (
        db.query(CaseModel)
        .filter(CaseModel.user_id == user_id)
        .order_by(CaseModel.created_at.desc())
        .all()
    )
    return cases


@router.post(
    "/workspace/documents",
    response_model=WorkspaceDocumentRead,
)
def create_workspace_document(
    payload: WorkspaceDocumentCreate,
    db: Session = Depends(get_db),
):
    """
    Создать документ из Workspace (редактор A4).
    """
    case = None
    if payload.case_id is not None:
        case = (
            db.query(CaseModel)
            .filter(
                CaseModel.id == payload.case_id,
                CaseModel.user_id == payload.user_id,
            )
            .first()
        )
        if case is None:
            raise HTTPException(status_code=404, detail="Дело не найдено")

    doc = DocumentModel(
        user_id=payload.user_id,
        case_id=payload.case_id,
        name=payload.name,
        status="draft",
        content_html=payload.content_html,
        file_path="PENDING",  # будет заменён при генерации DOCX/PDF
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    if case is not None:
        case.last_user_activity_at = datetime.utcnow()
        db.add(case)
        db.commit()

    return doc


@router.get(
    "/workspace/documents",
    response_model=List[WorkspaceDocumentRead],
)
def list_workspace_documents(
    user_id: int,
    db: Session = Depends(get_db),
):
    """
    Список документов пользователя для страницы 'Документы'.
    Показываем только те документы, у которых есть content_html (Workspace),
    чтобы не мешать DOCX-документы из шаблонов.
    """
    docs = (
        db.query(DocumentModel)
        .filter(
            DocumentModel.user_id == user_id,
            DocumentModel.content_html != None,  # только workspace-документы
        )
        .order_by(DocumentModel.created_at.desc())
        .all()
    )
    return docs


@router.patch(
    "/workspace/documents/{document_id}",
    response_model=WorkspaceDocumentRead,
)
def update_workspace_document(
    document_id: int,
    payload: WorkspaceDocumentUpdate,
    db: Session = Depends(get_db),
):
    """
    Обновление документа из Workspace.
    """
    doc = (
        db.query(DocumentModel)
        .filter(DocumentModel.id == document_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")

    if payload.name is not None:
        doc.name = payload.name
    if payload.content_html is not None:
        doc.content_html = payload.content_html
    if payload.status is not None:
        doc.status = payload.status

    doc.updated_at = datetime.utcnow()
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


@router.post(
    "/workspace/cases/{case_id}/attachments",
    response_model=AttachmentRead,
)
async def upload_case_attachment(
    case_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Загрузка вложения для конкретного дела.
    Сохраняем файл на диск и создаём запись в таблице attachments.
    """
    case = db.query(CaseModel).filter(CaseModel.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Дело не найдено")

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Empty file")
    # PDF pages limit (max 50)
    if (file.filename or "").lower().endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            pages = len(reader.pages)
        except Exception:
            raise HTTPException(status_code=400, detail="PDF could not be parsed or exceeds page limits. Please reduce the number of pages.")
        if pages > 50:
            raise HTTPException(
                status_code=400,
                detail=f"PDF contains {pages} pages. Maximum allowed is 50. Please reduce the number of pages."
            )


    if len(contents) > 25 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File too large",
        )

    backend_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    attach_dir = os.path.join(backend_root, "workspaces", "cases", str(case_id), "attachments")
    os.makedirs(attach_dir, exist_ok=True)

    original_name = (file.filename or "file").strip()
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", original_name) or "file"
    stored_name = f"{uuid.uuid4()}_{safe_name}"
    stored_path = os.path.join(attach_dir, stored_name)

    with open(stored_path, "wb") as f:
        f.write(contents)

    att = AttachmentModel(
        case_id=case_id,
        original_name=original_name,
        stored_path=stored_path,
    )
    db.add(att)
    db.commit()
    db.refresh(att)
    return att

@router.get(
    "/workspace/cases/{case_id}/attachments",
    response_model=List[AttachmentRead],
)
def list_case_attachments(
    case_id: int,
    db: Session = Depends(get_db),
):
    """
    Список вложений для конкретного дела.
    Загрузку файлов добавим отдельным шагом.
    """
    attachments = (
        db.query(AttachmentModel)
        .filter(AttachmentModel.case_id == case_id)
        .order_by(AttachmentModel.uploaded_at.desc())
        .all()
    )
    return attachments

@router.get("/workspace/attachments/{attachment_id}/text")
def get_attachment_text(
    attachment_id: int,
    db: Session = Depends(get_db),
):
    """Вернуть текст вложения (пока гарантированно поддерживаем .txt)."""
    att = db.query(AttachmentModel).filter(AttachmentModel.id == attachment_id).first()
    if not att:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Attachment not found")

    path = att.stored_path
    if not path or not os.path.isfile(path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Attachment file not found on disk")

    name = att.original_name or os.path.basename(path)
    ext = os.path.splitext(name)[1].lower()

    # 1) Точно поддерживаем текстовые файлы
    if ext in (".txt", ".md", ".csv", ".log"):
        try:
            text = Path(path).read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Failed to read text: {e}")
        return {
            "id": att.id,
            "case_id": att.case_id,
            "original_name": att.original_name,
            "stored_path": att.stored_path,
            "kind": "text",
            "text": text,
        }


    # 1.1) PDF (текстовый, без OCR)
    if ext == ".pdf":
        try:
            reader = PdfReader(path)
            pages_text = []
            for p in reader.pages:
                t = p.extract_text() or ""
                pages_text.append(t)
            text = "\n\n".join(pages_text)
        except Exception as e:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Failed to read PDF: {e}")
        return {
            "id": att.id,
            "case_id": att.case_id,
            "original_name": att.original_name,
            "stored_path": att.stored_path,
            "kind": "pdf",
            "text": text,
        }

    # 2) Остальные форматы — через общий экстрактор (PDF OCR / DOCX / IMG OCR)
    from legal_doc.text_extractor import extract_attachment_text
    text = extract_attachment_text(attachment_id)
    return {
        "id": att.id,
        "case_id": att.case_id,
        "original_name": att.original_name,
        "stored_path": att.stored_path,
        "kind": ext.lstrip("."),
        "text": text,
    }

