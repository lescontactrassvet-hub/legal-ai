import os
import uuid
from typing import Dict

from fastapi import UploadFile

# LEGALAI: безопасный импорт python-docx.
# В Termux python-docx может отсутствовать, на сервере Ubuntu он будет установлен.
try:
    from docx import Document
except ImportError:
    Document = None


BASE_DIR = os.path.dirname(__file__)
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
GENERATED_DIR = os.path.join(BASE_DIR, "generated_docs")


def ensure_dirs():
    """Ensure template and generated documents directories exist."""
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    os.makedirs(GENERATED_DIR, exist_ok=True)


async def save_template_file(upload_file: UploadFile, owner_id: int) -> str:
    """
    Save an uploaded template file to the templates directory.
    Returns the absolute file path.
    """
    ensure_dirs()
    filename = f"{owner_id}_{uuid.uuid4()}_{upload_file.filename}"
    file_path = os.path.join(TEMPLATE_DIR, filename)
    contents = await upload_file.read()
    with open(file_path, "wb") as f:
        f.write(contents)
    return file_path


def fill_template(template_file_path: str, fields: Dict[str, str]) -> str:
    """
    Fill the docx template with the given fields and return the path
    to the generated document.
    """
    ensure_dirs()

    # Если python-docx недоступен (Termux), явно сообщаем об этом.
    if Document is None:
        raise RuntimeError(
            "DOCX operations are unavailable in this environment "
            "(python-docx is not installed). Запусти backend на сервере."
        )

    doc = Document(template_file_path)

    # replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in fields.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in fields.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    new_filename = f"{uuid.uuid4()}.docx"
    new_file_path = os.path.join(GENERATED_DIR, new_filename)
    doc.save(new_file_path)
    return new_file_path

def get_status(attachment_id: int):
    # Временная безопасная реализация статуса
    return {
        "attachment_id": attachment_id,
        "used": False,
        "reason": "status not yet implemented"
    }
